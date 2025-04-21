from fastapi import APIRouter, HTTPException, Response
from pydantic import BaseModel
from typing import Dict, Any
import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

router = APIRouter()

class ResumeExportRequest(BaseModel):
    name: str = ""
    emails: list = []
    phones: list = []
    links: list = []
    education: list = []
    work_experience: list = []
    skills: list = []
    summary: str = ""

@router.post("/docx")
def export_docx(data: ResumeExportRequest):
    doc = Document()
    doc.add_heading(data.name or "Resume", 0)
    if data.summary:
        doc.add_paragraph(data.summary)
    if data.emails:
        doc.add_paragraph(f"Email: {', '.join(data.emails)}")
    if data.phones:
        doc.add_paragraph(f"Phone: {', '.join(data.phones)}")
    if data.links:
        doc.add_paragraph(f"Links: {', '.join(data.links)}")
    if data.skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data.skills))
    if data.education:
        doc.add_heading("Education", level=1)
        for edu in data.education:
            doc.add_paragraph(edu)
    if data.work_experience:
        doc.add_heading("Work Experience", level=1)
        for work in data.work_experience:
            doc.add_paragraph(work)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(content=buf.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=resume.docx"})

@router.post("/pdf")
def export_pdf(data: ResumeExportRequest):
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 40
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, data.name or "Resume")
    y -= 30
    c.setFont("Helvetica", 12)
    if data.summary:
        c.drawString(40, y, data.summary)
        y -= 20
    if data.emails:
        c.drawString(40, y, f"Email: {', '.join(data.emails)}")
        y -= 20
    if data.phones:
        c.drawString(40, y, f"Phone: {', '.join(data.phones)}")
        y -= 20
    if data.links:
        c.drawString(40, y, f"Links: {', '.join(data.links)}")
        y -= 20
    if data.skills:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y, "Skills:")
        y -= 18
        c.setFont("Helvetica", 12)
        c.drawString(60, y, ", ".join(data.skills))
        y -= 20
    if data.education:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y, "Education:")
        y -= 18
        c.setFont("Helvetica", 12)
        for edu in data.education:
            c.drawString(60, y, edu)
            y -= 16
    if data.work_experience:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y, "Work Experience:")
        y -= 18
        c.setFont("Helvetica", 12)
        for work in data.work_experience:
            c.drawString(60, y, work)
            y -= 16
    c.save()
    buf.seek(0)
    return Response(content=buf.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=resume.pdf"})
