"""
ResumeSense 2.0 - Parser Service
Robust extraction of text from PDF, DOCX, and TXT files.
"""
import re
from pathlib import Path
from typing import Optional

# PDF extraction
try:
    import fitz  # PyMuPDF - more reliable than pdfminer
except ImportError:
    fitz = None

from pdfminer.high_level import extract_text as pdfminer_extract

# DOCX extraction  
from docx import Document


class ParserService:
    """Handles document parsing and text extraction."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
    
    @classmethod
    def extract_text(cls, file_path: str | Path) -> str:
        """
        Extract text content from a file.
        
        Args:
            file_path: Path to the file to parse.
            
        Returns:
            Extracted text content.
            
        Raises:
            ValueError: If file type is not supported.
            FileNotFoundError: If file does not exist.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        ext = path.suffix.lower()
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {cls.SUPPORTED_EXTENSIONS}")
        
        if ext == ".pdf":
            return cls._extract_pdf(path)
        elif ext in {".docx", ".doc"}:
            return cls._extract_docx(path)
        elif ext == ".txt":
            return cls._extract_txt(path)
        
        return ""
    
    @classmethod
    def _extract_pdf(cls, path: Path) -> str:
        """Extract text from PDF using PyMuPDF (preferred) or pdfminer (fallback)."""
        text = ""
        
        # Try PyMuPDF first (faster and more reliable)
        if fitz:
            try:
                doc = fitz.open(str(path))
                for page in doc:
                    text += page.get_text()
                doc.close()
                if text.strip():
                    return cls._clean_text(text)
            except Exception:
                pass  # Fall back to pdfminer
        
        # Fallback to pdfminer
        try:
            text = pdfminer_extract(str(path))
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {e}")
        
        return cls._clean_text(text)
    
    @classmethod
    def _extract_docx(cls, path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return cls._clean_text("\n".join(paragraphs))
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {e}")
    
    @classmethod
    def _extract_txt(cls, path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return cls._clean_text(f.read())
        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {e}")
    
    @classmethod
    def _clean_text(cls, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove null bytes and other control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Collapse multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


# Convenience function
def parse_resume(file_path: str | Path) -> str:
    """Parse a resume file and return extracted text."""
    return ParserService.extract_text(file_path)
